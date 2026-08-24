# -*- coding: utf-8 -*-
"""
word_report.py — 불량률 분석 Word 보고서 생성
실제 세션 데이터(raw_rows, cache)를 받아 .docx 바이트를 반환
"""
import io, os
from collections import defaultdict
from datetime import datetime

try:
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    MPL_OK = True
except ImportError:
    MPL_OK = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

TOP_N = 7  # 상위 N개 불량 유형 표시, 나머지는 기타

# ── 색상 ─────────────────────────────────────────────────────────
C_HEADER = RGBColor(0x1A, 0x35, 0x57) if DOCX_OK else None
C_SUB    = RGBColor(0x2E, 0x6D, 0xA4) if DOCX_OK else None
C_RED    = RGBColor(0xC0, 0x39, 0x2B) if DOCX_OK else None
C_GREEN  = RGBColor(0x27, 0xAE, 0x60) if DOCX_OK else None
C_DARK   = RGBColor(0x22, 0x22, 0x22) if DOCX_OK else None
C_THEAD  = "1A3557"
C_TBODY  = "D5E8F0"
C_WHITE  = "FFFFFF"

MN = "#1A3557"; MB = "#2E6DA4"; ML = "#7FB3D3"
MR = "#C0392B"; MO = "#E67E22"; MG = "#27AE60"
PALETTE = [MN, MB, ML, MR, MO, MG, "#8E44AD", "#16A085", "#D35400", "#2C3E50"]

# ── 한글 폰트 설정 ────────────────────────────────────────────────
import os
WORD_FONT = "맑은 고딕"
_FONT_PATHS = [
    r"C:\Windows\Fonts\malgun.ttf",          # Windows 맑은 고딕
    r"C:\Windows\Fonts\gulim.ttc",           # Windows 굴림
    r"C:\Windows\Fonts\batang.ttc",          # Windows 바탕
    "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",          # Linux Nanum
    "/usr/share/fonts/truetype/nanum/NanumBarunGothic.ttf",
    "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",   # Noto CJK
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
]
_MPL_FONT = None
if MPL_OK:
    for _fp in _FONT_PATHS:
        if os.path.exists(_fp):
            try:
                fm.fontManager.addfont(_fp)
                _MPL_FONT = fm.FontProperties(fname=_fp).get_name()
                break
            except Exception:
                pass
    if _MPL_FONT:
        plt.rcParams['font.family'] = _MPL_FONT
    else:
        # 시스템에 등록된 한글 가능 폰트 탐색
        for _fn in ['Malgun Gothic', 'NanumGothic', 'NanumBarunGothic',
                    'Apple SD Gothic Neo', 'AppleGothic']:
            try:
                fm.findfont(fm.FontProperties(family=_fn), fallback_to_default=False)
                plt.rcParams['font.family'] = _fn
                _MPL_FONT = _fn
                break
            except Exception:
                pass
    plt.rcParams['axes.unicode_minus'] = False


# ── 데이터 집계 함수 ──────────────────────────────────────────────

def _safe_int(v):
    try:
        return int(v or 0)
    except (TypeError, ValueError):
        return 0


def _rate(defect, inspec):
    return round(defect / inspec * 100, 2) if inspec > 0 else 0.0


def _top5_others(counter: dict):
    """상위 TOP_N + 기타로 묶어서 반환"""
    sorted_items = sorted(counter.items(), key=lambda x: x[1], reverse=True)
    top = sorted_items[:TOP_N]
    others_total = sum(v for _, v in sorted_items[TOP_N:])
    if others_total > 0:
        top.append(("기타", others_total))
    return top  # [(name, count), ...]


def aggregate(raw_rows: list, cache: dict) -> dict:
    """세션 데이터에서 보고서용 집계 수행"""

    # 불량명칭 → 표준명칭 매핑 (cache: {defect_raw: [(part, std, sc, meth, rev, note)]})
    def get_std(defect_raw):
        results = cache.get(defect_raw, [])
        if not results:
            return defect_raw
        return results[0][1] or defect_raw  # 첫 번째 결과의 std

    # 전체 기간 파악
    dates = [r.get('date', '') for r in raw_rows if r.get('date')]
    period_start = min(dates) if dates else ''
    period_end   = max(dates) if dates else ''
    if period_start and period_end:
        period = f"{period_start[:7]} ~ {period_end[:7]}"
    else:
        period = datetime.now().strftime('%Y-%m') + " 기준"

    # ── report_no 기준 중복 제거 사전 구축 ────────────────────────
    # 불량상세 시트는 불량유형별로 행이 분리되어 검사수량/최종불합격/2차검사가 반복됨.
    # report_no 단위로 한 번만 합산해야 정확한 불량률이 계산됨.
    _seen_rno: dict = {}  # {(group_key, report_no): True}
    def _add_inspec(seen_key, r, inspec_d, final_d, second_d, group, defect_d=None):
        """report_no 기준 dedup — inspec/final/second/qty_1st 각 1회만 합산"""
        rno = str(r.get('report_no') or '').strip() or f'__row_{id(r)}'
        key = (group, rno)
        if key not in seen_key:
            seen_key[key] = True
            inspec_d[group] += _safe_int(r.get('inspec'))
            final_d[group]  += _safe_int(r.get('최종불합격수량'))
            second_d[group] += _safe_int(r.get('2차검사수량'))
            if defect_d is not None:
                defect_d[group] += _safe_int(r.get('qty_1st'))

    # ── 1. 전체 요약 ───────────────────────────────
    _seen1: dict = {}
    _tot_inspec = defaultdict(int); _tot_final = defaultdict(int)
    _tot_second = defaultdict(int); _tot_defect = defaultdict(int)
    for r in raw_rows:
        _add_inspec(_seen1, r, _tot_inspec, _tot_final, _tot_second, 'all', _tot_defect)
    total_inspec        = _tot_inspec['all']
    total_defect        = _tot_defect['all']
    total_final_defect  = _tot_final['all']
    total_second_inspec = _tot_second['all']
    total_rate       = _rate(total_defect, total_inspec)
    total_final_rate = _rate(total_final_defect, total_inspec)
    total_correction = (_rate(total_second_inspec - total_final_defect, total_second_inspec)
                        if total_second_inspec > 0 else 0.0)

    # ── 2. 월별 ────────────────────────────────────
    _seen2: dict = {}
    monthly_inspec  = defaultdict(int)
    monthly_defect  = defaultdict(int)
    monthly_final   = defaultdict(int)
    monthly_second  = defaultdict(int)
    for r in raw_rows:
        ym = (r.get('date') or '')[:7]
        if not ym:
            continue
        _add_inspec(_seen2, r, monthly_inspec, monthly_final, monthly_second, ym, monthly_defect)
    monthly = sorted([
        {"month": ym, "inspec": monthly_inspec[ym],
         "defect": monthly_defect[ym],
         "rate": _rate(monthly_defect[ym], monthly_inspec[ym]),
         "final_defect": monthly_final[ym],
         "final_rate": _rate(monthly_final[ym], monthly_inspec[ym]),
         "second_inspec": monthly_second[ym],
         "correction_rate": (_rate(monthly_second[ym] - monthly_final[ym], monthly_second[ym])
                             if monthly_second[ym] > 0 else 0.0)}
        for ym in monthly_inspec
    ], key=lambda x: x["month"])

    # ── 3. 업체별 ──────────────────────────────────
    _seen3: dict = {}
    client_inspec  = defaultdict(int)
    client_defect  = defaultdict(int)
    client_final   = defaultdict(int)
    client_second  = defaultdict(int)
    for r in raw_rows:
        c = str(r.get('client') or r.get('buyer') or '미확인').strip()
        _add_inspec(_seen3, r, client_inspec, client_final, client_second, c, client_defect)
    by_client = sorted([
        {"name": c, "inspec": client_inspec[c],
         "defect": client_defect[c],
         "rate": _rate(client_defect[c], client_inspec[c]),
         "final_defect": client_final[c],
         "final_rate": _rate(client_final[c], client_inspec[c]),
         "second_inspec": client_second[c],
         "correction_rate": (_rate(client_second[c] - client_final[c], client_second[c])
                             if client_second[c] > 0 else 0.0)}
        for c in client_inspec
    ], key=lambda x: x["rate"], reverse=True)

    # ── 3-2. 국가별 불량률 ─────────────────────────
    _seen4: dict = {}
    country_inspec = defaultdict(int)
    country_defect = defaultdict(int)
    country_final  = defaultdict(int)
    country_second = defaultdict(int)
    for r in raw_rows:
        country = str(r.get('region1') or '미확인').strip() or '미확인'
        _add_inspec(_seen4, r, country_inspec, country_final, country_second, country, country_defect)
    by_country = sorted([
        {"name": c, "inspec": country_inspec[c],
         "defect": country_defect[c],
         "rate": _rate(country_defect[c], country_inspec[c]),
         "final_defect": country_final[c],
         "final_rate": _rate(country_final[c], country_inspec[c]),
         "second_inspec": country_second[c],
         "correction_rate": (_rate(country_second[c] - country_final[c], country_second[c])
                             if country_second[c] > 0 else 0.0)}
        for c in country_inspec
    ], key=lambda x: x["rate"], reverse=True)

    # ── 4. 전체 세부 불량 유형 ─────────────────────
    defect_count = defaultdict(int)
    for r in raw_rows:
        std = get_std(r.get('defect_raw', ''))
        defect_count[std] += _safe_int(r.get('qty_total')) or 1
    total_defect_cnt = sum(defect_count.values())
    defect_top = _top5_others(defect_count)
    defect_types = [
        {"type": name, "count": cnt,
         "pct": round(cnt / total_defect_cnt * 100, 1) if total_defect_cnt else 0}
        for name, cnt in defect_top
    ]

    # 상위 유형 이름 목록 (교차표 컬럼용)
    top_type_names = [d["type"] for d in defect_types]

    # ── 5. 업체별 세부 불량 유형 ───────────────────
    client_type = defaultdict(lambda: defaultdict(int))
    for r in raw_rows:
        c = str(r.get('client') or r.get('buyer') or '미확인').strip()
        std = get_std(r.get('defect_raw', ''))
        cnt = _safe_int(r.get('qty_total')) or 1
        if std in top_type_names:
            client_type[c][std] += cnt
        else:
            client_type[c]["기타"] += cnt
    client_defect_table = {
        c: {t: client_type[c].get(t, 0) for t in top_type_names}
        for c in client_inspec
    }

    # ── 6. 공장별 불량률 ───────────────────────────
    _seen6: dict = {}
    _fac_final = defaultdict(int); _fac_second = defaultdict(int)
    factory_inspec = defaultdict(int)
    factory_defect = defaultdict(int)
    for r in raw_rows:
        f = str(r.get('factory') or '미확인').strip()
        _add_inspec(_seen6, r, factory_inspec, _fac_final, _fac_second, f, factory_defect)
    by_factory = sorted([
        {"name": f, "inspec": factory_inspec[f],
         "defect": factory_defect[f],
         "rate": _rate(factory_defect[f], factory_inspec[f])}
        for f in factory_inspec
    ], key=lambda x: x["rate"], reverse=True)

    # ── 7. 공장별 세부 불량 유형 ───────────────────
    factory_type = defaultdict(lambda: defaultdict(int))
    for r in raw_rows:
        f = str(r.get('factory') or '미확인').strip()
        std = get_std(r.get('defect_raw', ''))
        cnt = _safe_int(r.get('qty_total')) or 1
        if std in top_type_names:
            factory_type[f][std] += cnt
        else:
            factory_type[f]["기타"] += cnt
    factory_defect_table = {
        f["name"]: {t: factory_type[f["name"]].get(t, 0) for t in top_type_names}
        for f in by_factory
    }

    # ── 8. 품목 유형별 비율 ────────────────────────
    _신발_KW = ['신발', 'SHOE', 'SHOES', 'SNEAKER', 'BOOT', 'BOOTS', 'SANDAL',
                '부츠', '샌들', '스니커', 'SLIPPER', '슬리퍼']
    _잡화_KW = ['가방', 'BAG', '지갑', 'WALLET', '파우치', 'POUCH',
                '모자', 'HAT', 'CAP', '머플러', 'MUFFLER', 'SCARF',
                '장갑', 'GLOVE', '벨트', 'BELT', '우산', '백팩', 'BACKPACK']
    def _classify(item_str):
        s = str(item_str or '').upper()
        for kw in _신발_KW:
            if kw in s: return '신발'
        for kw in _잡화_KW:
            if kw in s: return '잡화'
        return '의류'
    _seen8: dict = {}
    _typ_final = defaultdict(int); _typ_second = defaultdict(int)
    type_inspec = defaultdict(int)
    type_defect = defaultdict(int)
    for r in raw_rows:
        ptype = r.get('product_type') or _classify(r.get('item', ''))
        _add_inspec(_seen8, r, type_inspec, _typ_final, _typ_second, ptype, type_defect)
    by_item_type = [
        {"name": t, "inspec": type_inspec[t],
         "defect": type_defect[t],
         "rate": _rate(type_defect[t], type_inspec[t])}
        for t in ['의류', '잡화', '신발'] if type_inspec[t] > 0
    ]

    return {
        "period": period,
        "summary": {
            "inspec": total_inspec, "defect": total_defect, "rate": total_rate,
            "final_defect": total_final_defect, "final_rate": total_final_rate,
            "second_inspec": total_second_inspec, "correction": total_correction,
        },
        "monthly": monthly,
        "by_country": by_country,
        "by_client": by_client,
        "by_item_type": by_item_type,
        "defect_types": defect_types,
        "top_type_names": top_type_names,
        "client_defect_table": client_defect_table,
        "by_factory": by_factory,
        "factory_defect_table": factory_defect_table,
    }


# ── 차트 생성 ─────────────────────────────────────────────────────

def _fig_bytes(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=150, bbox_inches='tight',
                facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)
    return buf


def _chart_item_type(by_item_type):
    """품목 유형별(의류/잡화/신발) 파이차트 + 불량률 바차트"""
    if not by_item_type:
        return None
    names  = [d["name"]  for d in by_item_type]
    inspec = [d["inspec"] for d in by_item_type]
    rates  = [d["rate"]  for d in by_item_type]
    colors_pie = [MB, MR, '#F4A261'][:len(names)]

    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 3.5))

    wedges, texts, autotexts = ax1.pie(
        inspec, labels=names, autopct='%1.1f%%',
        colors=colors_pie, startangle=90,
        textprops={'fontsize': 9},
    )
    ax1.set_title('품목 유형별 검사 비율', fontsize=10, fontweight='bold')

    bar_colors = [MB, MR, '#F4A261'][:len(names)]
    bars = ax2.bar(names, rates, color=bar_colors, width=0.4)
    for b, r in zip(bars, rates):
        ax2.text(b.get_x() + b.get_width()/2, b.get_height() + 0.03,
                 f'{r:.2f}%', ha='center', va='bottom', fontsize=9)
    ax2.set_ylabel('불량률 (%)', fontsize=9)
    ax2.set_title('품목 유형별 불량률', fontsize=10, fontweight='bold')
    ax2.set_ylim(0, max(rates)*1.4 if rates else 10)
    ax2.tick_params(axis='x', labelsize=9)
    ax2.grid(axis='y', alpha=0.3)
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)

    fig.tight_layout()
    return _fig_bytes(fig)


def _chart_country(by_country, avg_rate):
    names = [d["name"] for d in by_country]
    rates = [d["rate"] for d in by_country]
    colors = [MR if r > avg_rate else MB for r in rates]
    fig, ax = plt.subplots(figsize=(max(6, len(names)*1.2), 3))
    bars = ax.bar(names, rates, color=colors, width=0.5)
    if avg_rate:
        ax.axhline(avg_rate, color=MN, linewidth=1.5, linestyle='--',
                   label=f'평균 {avg_rate:.2f}%')
    for b, r in zip(bars, rates):
        ax.text(b.get_x() + b.get_width()/2, b.get_height() + 0.04,
                f'{r:.2f}%', ha='center', va='bottom', fontsize=9)
    ax.set_ylim(0, max(rates) * 1.3 if rates else 10)
    ax.set_ylabel('불량률 (%)', fontsize=9)
    ax.set_title('국가별 불량률', fontsize=11, fontweight='bold')
    ax.tick_params(axis='x', labelsize=9)
    ax.legend(fontsize=8)
    ax.grid(axis='y', alpha=0.3)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return _fig_bytes(fig)


def _chart_monthly(monthly, avg_rate):
    months = [d["month"] for d in monthly]
    rates  = [d["rate"] for d in monthly]
    fig, ax = plt.subplots(figsize=(8, 3))
    bars = ax.bar(months, rates, color=MB, width=0.5, zorder=2)
    if avg_rate:
        ax.axhline(avg_rate, color=MR, linewidth=1.5, linestyle='--',
                   label=f'평균 {avg_rate:.2f}%')
    for b, r in zip(bars, rates):
        ax.text(b.get_x() + b.get_width()/2, b.get_height() + 0.05,
                f'{r:.2f}%', ha='center', va='bottom', fontsize=8)
    ax.set_ylim(0, max(rates) * 1.3 if rates else 10)
    ax.set_ylabel('불량률 (%)', fontsize=9)
    ax.set_title('월별 불량률 추이', fontsize=11, fontweight='bold')
    ax.tick_params(axis='x', labelsize=8, rotation=20)
    ax.legend(fontsize=8)
    ax.grid(axis='y', alpha=0.3, zorder=1)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return _fig_bytes(fig)


def _chart_client(by_client, avg_rate):
    names = [d["name"] for d in by_client]
    rates = [d["rate"] for d in by_client]
    colors = [MR if r > avg_rate else MB for r in rates]
    fig, ax = plt.subplots(figsize=(7, max(3, len(names) * 0.5 + 1)))
    bars = ax.barh(names, rates, color=colors, height=0.5)
    if avg_rate:
        ax.axvline(avg_rate, color=MN, linewidth=1.5, linestyle='--',
                   label=f'평균 {avg_rate:.2f}%')
    for b, r in zip(bars, rates):
        ax.text(r + 0.05, b.get_y() + b.get_height()/2,
                f'{r:.2f}%', va='center', fontsize=8)
    ax.set_xlim(0, max(rates) * 1.3 if rates else 10)
    ax.set_xlabel('불량률 (%)', fontsize=9)
    ax.set_title('업체별 불량률', fontsize=11, fontweight='bold')
    ax.tick_params(axis='y', labelsize=8)
    ax.legend(fontsize=8)
    ax.grid(axis='x', alpha=0.3)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return _fig_bytes(fig)


def _chart_defect(defect_types):
    labels = [f"#{i+1} {d['type']}" for i, d in enumerate(defect_types)]
    counts = [d["count"] for d in defect_types]
    pal = PALETTE[:len(counts)]
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9, 3.5))
    ax1.pie(counts, labels=labels, colors=pal, autopct='%1.1f%%',
            startangle=90, textprops={'fontsize': 7}, pctdistance=0.82)
    ax1.set_title('불량 유형 분포', fontsize=10, fontweight='bold')
    x_idx = range(len(counts))
    ax2.bar(x_idx, counts, color=pal, width=0.55)
    for i, c in enumerate(counts):
        ax2.text(i, c + max(counts)*0.02, str(c), ha='center', va='bottom', fontsize=8)
    ax2.set_xticks(list(x_idx))
    ax2.set_xticklabels([f"#{i+1}" for i in range(len(counts))], fontsize=9)
    ax2.set_ylabel('건수', fontsize=9)
    ax2.set_title('불량 유형별 건수', fontsize=10, fontweight='bold')
    ax2.grid(axis='y', alpha=0.3)
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    fig.tight_layout(pad=2)
    return _fig_bytes(fig)


def _chart_factory(by_factory, avg_rate):
    names = [d["name"] for d in by_factory]
    rates = [d["rate"] for d in by_factory]
    colors = [MR if r > avg_rate else MB for r in rates]
    fig, ax = plt.subplots(figsize=(max(7, len(names)*1.2), 3))
    bars = ax.bar(names, rates, color=colors, width=0.5)
    if avg_rate:
        ax.axhline(avg_rate, color=MN, linewidth=1.5, linestyle='--',
                   label=f'평균 {avg_rate:.2f}%')
    for b, r in zip(bars, rates):
        ax.text(b.get_x() + b.get_width()/2, b.get_height() + 0.04,
                f'{r:.2f}%', ha='center', va='bottom', fontsize=8)
    ax.set_ylim(0, max(rates) * 1.3 if rates else 10)
    ax.set_ylabel('불량률 (%)', fontsize=9)
    ax.set_title('공장별 불량률', fontsize=11, fontweight='bold')
    ax.tick_params(axis='x', labelsize=8, rotation=15)
    ax.legend(fontsize=8)
    ax.grid(axis='y', alpha=0.3)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return _fig_bytes(fig)


# ── 1차 / 최종 이중 바차트 (수정합격률 없음) ──────────────────────

def _chart_monthly_dual(monthly, avg1, avg2):
    """1차 / 최종 불량률 월별 이중 바차트"""
    months = [d["month"] for d in monthly]
    rates1 = [d["rate"] for d in monthly]
    rates2 = [d.get("final_rate", 0) for d in monthly]
    x = list(range(len(months))); width = 0.36
    fig, ax = plt.subplots(figsize=(max(9, len(months) * 0.8), 3.5))
    b1 = ax.bar([i - width/2 for i in x], rates1, width, color=MB, label='1차 불량률', zorder=2)
    b2 = ax.bar([i + width/2 for i in x], rates2, width, color=MR, label='최종 불량률', zorder=2)
    if avg1:
        ax.axhline(avg1, color=MB, linewidth=1.2, linestyle='--', alpha=0.7,
                   label=f'1차 평균 {avg1:.2f}%')
    if avg2:
        ax.axhline(avg2, color=MR, linewidth=1.2, linestyle='--', alpha=0.7,
                   label=f'최종 평균 {avg2:.2f}%')
    for b, r in zip(b1.patches, rates1):
        if r > 0:
            ax.text(b.get_x() + b.get_width()/2, b.get_height() + 0.03,
                    f'{r:.2f}%', ha='center', va='bottom', fontsize=7, color=MB)
    for b, r in zip(b2.patches, rates2):
        if r > 0:
            ax.text(b.get_x() + b.get_width()/2, b.get_height() + 0.03,
                    f'{r:.2f}%', ha='center', va='bottom', fontsize=7, color=MR)
    ax.set_xticks(x); ax.set_xticklabels(months, rotation=20, fontsize=8)
    ax.set_ylabel('불량률 (%)', fontsize=9)
    ax.set_title('월별 불량률 추이  (1차 / 최종)', fontsize=11, fontweight='bold')
    ax.legend(fontsize=8, ncol=2); ax.grid(axis='y', alpha=0.3, zorder=1)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    fig.tight_layout(); return _fig_bytes(fig)


def _chart_country_dual(by_country, avg1, avg2):
    """국가별 1차 / 최종 이중 바차트"""
    names  = [d["name"] for d in by_country]
    rates1 = [d["rate"] for d in by_country]
    rates2 = [d.get("final_rate", 0) for d in by_country]
    x = list(range(len(names))); width = 0.36
    fig, ax = plt.subplots(figsize=(max(6, len(names) * 1.4), 3.5))
    ax.bar([i - width/2 for i in x], rates1, width, color=MB, label='1차 불량률')
    ax.bar([i + width/2 for i in x], rates2, width, color=MR, label='최종 불량률')
    if avg1:
        ax.axhline(avg1, color=MB, linewidth=1.2, linestyle='--', alpha=0.7,
                   label=f'1차 평균 {avg1:.2f}%')
    if avg2:
        ax.axhline(avg2, color=MR, linewidth=1.2, linestyle='--', alpha=0.7,
                   label=f'최종 평균 {avg2:.2f}%')
    ax.set_xticks(x); ax.set_xticklabels(names, fontsize=9)
    ax.set_ylabel('불량률 (%)', fontsize=9)
    ax.set_title('국가별 불량률  (1차 / 최종)', fontsize=11, fontweight='bold')
    ax.legend(fontsize=8, ncol=2); ax.grid(axis='y', alpha=0.3)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    fig.tight_layout(); return _fig_bytes(fig)


def _chart_client_dual(by_client, avg1, avg2):
    """업체별 1차 / 최종 수평 이중 바차트"""
    names  = [d["name"] for d in by_client]
    rates1 = [d["rate"] for d in by_client]
    rates2 = [d.get("final_rate", 0) for d in by_client]
    x = list(range(len(names))); width = 0.36
    fig, ax = plt.subplots(figsize=(7, max(3.5, len(names) * 0.6 + 1)))
    ax.barh([i + width/2 for i in x], rates1, width, color=MB, label='1차 불량률')
    ax.barh([i - width/2 for i in x], rates2, width, color=MR, label='최종 불량률')
    if avg1:
        ax.axvline(avg1, color=MB, linewidth=1.2, linestyle='--', alpha=0.7,
                   label=f'1차 평균 {avg1:.2f}%')
    if avg2:
        ax.axvline(avg2, color=MR, linewidth=1.2, linestyle='--', alpha=0.7,
                   label=f'최종 평균 {avg2:.2f}%')
    ax.set_yticks(x); ax.set_yticklabels(names, fontsize=8)
    ax.set_xlabel('불량률 (%)', fontsize=9)
    ax.set_title('업체별 불량률  (1차 / 최종)', fontsize=11, fontweight='bold')
    ax.legend(fontsize=8, ncol=2); ax.grid(axis='x', alpha=0.3)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    fig.tight_layout(); return _fig_bytes(fig)


# ── 전체 모드: 1차 / 최종 바 + 수정합격률 라인 (twin y-axis) ──────

def _chart_monthly_all(monthly, avg1, avg2, avg_corr):
    """전체 모드 - 월별 1차/최종 바 + 수정합격률 꺾은선 (우측 축)"""
    months = [d["month"] for d in monthly]
    rates1 = [d["rate"] for d in monthly]
    rates2 = [d.get("final_rate", 0) for d in monthly]
    corr   = [d.get("correction_rate", 0) for d in monthly]
    x = list(range(len(months))); width = 0.3
    fig, ax = plt.subplots(figsize=(max(9, len(months) * 0.9), 4))
    b1 = ax.bar([i - width for i in x], rates1, width, color=MB, label='1차 불량률', zorder=2)
    b2 = ax.bar([i         for i in x], rates2, width, color=MR, label='최종 불량률', zorder=2)
    if avg1: ax.axhline(avg1, color=MB, lw=1, ls='--', alpha=0.6, label=f'1차 평균 {avg1:.2f}%')
    if avg2: ax.axhline(avg2, color=MR, lw=1, ls='--', alpha=0.6, label=f'최종 평균 {avg2:.2f}%')
    for b, r in zip(b1.patches, rates1):
        if r > 0:
            ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.03,
                    f'{r:.2f}%', ha='center', va='bottom', fontsize=6.5, color=MB)
    for b, r in zip(b2.patches, rates2):
        if r > 0:
            ax.text(b.get_x()+b.get_width()/2, b.get_height()+0.03,
                    f'{r:.2f}%', ha='center', va='bottom', fontsize=6.5, color=MR)
    ax2 = ax.twinx()
    ax2.plot(x, corr, color=MG, lw=2, marker='o', ms=5, label='수정합격률', zorder=3)
    if avg_corr: ax2.axhline(avg_corr, color=MG, lw=1, ls='--', alpha=0.6)
    ax2.set_ylabel('수정합격률 (%)', fontsize=9, color=MG)
    ax2.tick_params(axis='y', labelcolor=MG, labelsize=8)
    ax2.set_ylim(0, 130)
    ax.set_xticks(x); ax.set_xticklabels(months, rotation=20, fontsize=8)
    ax.set_ylabel('불량률 (%)', fontsize=9); ax.set_ylim(bottom=0)
    ax.set_title('월별 불량률 추이  (1차 / 최종 / 수정합격률)', fontsize=11, fontweight='bold')
    lines1, lbl1 = ax.get_legend_handles_labels()
    lines2, lbl2 = ax2.get_legend_handles_labels()
    ax.legend(lines1+lines2, lbl1+lbl2, fontsize=7.5, ncol=3)
    ax.grid(axis='y', alpha=0.3, zorder=1)
    ax.spines['top'].set_visible(False)
    fig.tight_layout(); return _fig_bytes(fig)


def _chart_country_all(by_country, avg1, avg2, avg_corr):
    """전체 모드 - 국가별 1차/최종 바 + 수정합격률 라인"""
    names  = [d["name"] for d in by_country]
    rates1 = [d["rate"] for d in by_country]
    rates2 = [d.get("final_rate", 0) for d in by_country]
    corr   = [d.get("correction_rate", 0) for d in by_country]
    x = list(range(len(names))); width = 0.28
    fig, ax = plt.subplots(figsize=(max(6, len(names)*1.5), 3.8))
    ax.bar([i - width for i in x], rates1, width, color=MB, label='1차 불량률')
    ax.bar([i         for i in x], rates2, width, color=MR, label='최종 불량률')
    if avg1: ax.axhline(avg1, color=MB, lw=1, ls='--', alpha=0.6, label=f'1차 평균 {avg1:.2f}%')
    if avg2: ax.axhline(avg2, color=MR, lw=1, ls='--', alpha=0.6, label=f'최종 평균 {avg2:.2f}%')
    ax2 = ax.twinx()
    ax2.plot(x, corr, color=MG, lw=2, marker='o', ms=6, label='수정합격률', zorder=3)
    if avg_corr: ax2.axhline(avg_corr, color=MG, lw=1, ls='--', alpha=0.6)
    ax2.set_ylabel('수정합격률 (%)', fontsize=9, color=MG)
    ax2.tick_params(axis='y', labelcolor=MG, labelsize=8)
    ax2.set_ylim(0, 130)
    ax.set_xticks(x); ax.set_xticklabels(names, fontsize=9)
    ax.set_ylabel('불량률 (%)', fontsize=9); ax.set_ylim(bottom=0)
    ax.set_title('국가별 불량률  (1차 / 최종 / 수정합격률)', fontsize=11, fontweight='bold')
    lines1, lbl1 = ax.get_legend_handles_labels()
    lines2, lbl2 = ax2.get_legend_handles_labels()
    ax.legend(lines1+lines2, lbl1+lbl2, fontsize=7.5, ncol=3)
    ax.grid(axis='y', alpha=0.3)
    ax.spines['top'].set_visible(False)
    fig.tight_layout(); return _fig_bytes(fig)


def _chart_client_all(by_client, avg1, avg2, avg_corr):
    """전체 모드 - 업체별 1차/최종 수평 바 + 수정합격률 산점도"""
    names  = [d["name"] for d in by_client]
    rates1 = [d["rate"] for d in by_client]
    rates2 = [d.get("final_rate", 0) for d in by_client]
    corr   = [d.get("correction_rate", 0) for d in by_client]
    x = list(range(len(names))); width = 0.3
    fig, ax = plt.subplots(figsize=(8, max(3.5, len(names)*0.6+1)))
    ax.barh([i+width/2 for i in x], rates1, width, color=MB, label='1차 불량률')
    ax.barh([i-width/2 for i in x], rates2, width, color=MR, label='최종 불량률')
    if avg1: ax.axvline(avg1, color=MB, lw=1, ls='--', alpha=0.6, label=f'1차 평균 {avg1:.2f}%')
    if avg2: ax.axvline(avg2, color=MR, lw=1, ls='--', alpha=0.6, label=f'최종 평균 {avg2:.2f}%')
    ax2 = ax.twiny()
    ax2.scatter(corr, x, color=MG, s=60, marker='D', label='수정합격률', zorder=3)
    if avg_corr: ax2.axvline(avg_corr, color=MG, lw=1, ls='--', alpha=0.6)
    ax2.set_xlabel('수정합격률 (%)', fontsize=9, color=MG)
    ax2.tick_params(axis='x', labelcolor=MG, labelsize=8)
    ax2.set_xlim(0, 130)
    ax.set_yticks(x); ax.set_yticklabels(names, fontsize=8)
    ax.set_xlabel('불량률 (%)', fontsize=9); ax.set_xlim(left=0)
    ax.set_title('업체별 불량률  (1차 / 최종 / 수정합격률)', fontsize=11, fontweight='bold')
    lines1, lbl1 = ax.get_legend_handles_labels()
    lines2, lbl2 = ax2.get_legend_handles_labels()
    ax.legend(lines1+lines2, lbl1+lbl2, fontsize=7.5, ncol=3)
    ax.grid(axis='x', alpha=0.3)
    ax.spines['top'].set_visible(False)
    fig.tight_layout(); return _fig_bytes(fig)


# ── docx 헬퍼 ────────────────────────────────────────────────────

def _shd(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto')
    s.set(qn('w:fill'), hex_color); tcPr.append(s)


def _bdr(cell):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    b = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0'); el.set(qn('w:color'), 'AAAAAA')
        b.append(el)
    tcPr.append(b)


def _run(p, text, sz=11, bold=False, color=None):
    r = p.add_run(text)
    r.font.name = WORD_FONT; r.font.size = Pt(sz); r.font.bold = bold
    if color:
        r.font.color.rgb = color
    rPr = r._r.get_or_add_rPr()
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:eastAsia'), WORD_FONT); rPr.insert(0, rf)
    return r


def _ct(cell, text, bold=False, sz=9.5, color=None,
        align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    _run(p, str(text), sz=sz, bold=bold, color=color)


def _sec_title(doc, text, note=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(0 if note else 4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(f"■ {text}")
    r.font.name = WORD_FONT; r.font.size = Pt(13); r.font.bold = True
    r.font.color.rgb = C_HEADER
    rPr = r._r.get_or_add_rPr()
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:eastAsia'), WORD_FONT); rPr.insert(0, rf)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '6'); bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1A3557'); pBdr.append(bot); pPr.append(pBdr)
    if note:
        pn = doc.add_paragraph()
        pn.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pn.paragraph_format.space_before = Pt(1)
        pn.paragraph_format.space_after  = Pt(4)
        pn.paragraph_format.keep_with_next = True
        _run(pn, note, sz=10, color=RGBColor(0x88, 0x88, 0x88))


def _img_para(doc, img_bytes, width_cm=15):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.add_run().add_picture(img_bytes, width=Cm(width_cm))


def _spacer(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(2)


_LOGO_PATH = os.path.join(os.path.dirname(__file__), '..', 'BS 1-06 시그니처_국영문_가로형.png')


def _logo_header(doc, title, subtitle):
    """FITI CI 헤더: 제목(왼쪽) + 로고(오른쪽), 하단 파란 구분선"""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # 테이블 전체 테두리 제거
    tbl_pr = tbl._tbl.get_or_add_tblPr()
    tbl_bdr = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'), 'nil'); tbl_bdr.append(el)
    tbl_pr.append(tbl_bdr)

    lc, rc = tbl.cell(0, 0), tbl.cell(0, 1)

    # 왼쪽: 제목 + 부제
    lc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p1 = lc.paragraphs[0]
    p1.paragraph_format.space_before = Pt(4); p1.paragraph_format.space_after = Pt(0)
    _run(p1, title, sz=20, bold=True, color=C_HEADER)
    p2 = lc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(4)
    _run(p2, subtitle, sz=10, color=C_SUB)

    # 오른쪽: 로고
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p3 = rc.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p3.paragraph_format.space_before = Pt(4); p3.paragraph_format.space_after = Pt(4)
    if os.path.exists(_LOGO_PATH):
        p3.add_run().add_picture(_LOGO_PATH, width=Cm(5.5))

    # 셀 테두리: 하단만 파란 선
    def _cell_bdr(cell):
        tcPr = cell._tc.get_or_add_tcPr()
        b = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'right'):
            el = OxmlElement(f'w:{edge}'); el.set(qn('w:val'), 'nil'); b.append(el)
        bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '14'); bot.set(qn('w:space'), '0')
        bot.set(qn('w:color'), '003f85'); b.append(bot); tcPr.append(b)

    _cell_bdr(lc); _cell_bdr(rc)


def _hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr(); pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom'); bot.set(qn('w:val'), 'single')
    bot.set(qn('w:sz'), '12'); bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1A3557'); pBdr.append(bot); pPr.append(pBdr)


def _make_table(doc, headers, rows_data, col_widths=None):
    """헤더 + 데이터 행 테이블 생성
    rows_data: list of list[(text, bold, color, align)]
    """
    n_cols = len(headers)
    tbl = doc.add_table(rows=1 + len(rows_data), cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 헤더
    for cell, h in zip(tbl.rows[0].cells, headers):
        _ct(cell, h, bold=True, sz=9, color=RGBColor(0xFF,0xFF,0xFF))
        _shd(cell, C_THEAD); _bdr(cell)
    # 데이터
    for i, row_vals in enumerate(rows_data):
        row = tbl.rows[i + 1]
        bg = C_TBODY if (i % 2 == 1) else C_WHITE
        for cell, val in zip(row.cells, row_vals):
            if isinstance(val, tuple):
                text, bold, color, align = val
            else:
                text, bold, color, align = str(val), False, None, WD_ALIGN_PARAGRAPH.CENTER
            _ct(cell, text, bold=bold, sz=9.5, color=color, align=align)
            _shd(cell, bg); _bdr(cell)
    return tbl


# ── 보고서 생성 메인 ──────────────────────────────────────────────

def generate_word_report(raw_rows: list, cache: dict, orientation: str = 'portrait',
                         report_mode: str = '전체') -> bytes:
    """
    raw_rows + cache → .docx 파일 bytes 반환
    orientation:  'portrait'(세로형, 기본) | 'landscape'(가로형)
    report_mode:  '전체' | '1차 불량률' | '1차 불량률 + 최종 불량률'
    """
    if not DOCX_OK:
        raise RuntimeError("python-docx 설치 필요: pip install python-docx")

    data = aggregate(raw_rows, cache)
    period  = data["period"]
    summary = data["summary"]
    avg     = summary["rate"]
    IS_ALL  = (report_mode == '전체')                         # 1차+최종+수정합격률
    IS_DUAL = (report_mode == '1차 불량률 + 최종 불량률')     # 1차+최종만 (수정합격률 없음)
    IS_1ST  = (report_mode == '1차 불량률')
    IS_ANY_DUAL = IS_ALL or IS_DUAL                           # 1차+최종 표시 공통 조건
    avg_final = summary.get("final_rate", 0.0)
    avg_corr  = summary.get("correction", 0.0)

    # 방향별 페이지 설정
    IS_LAND = (orientation == 'landscape')
    _pw  = Cm(29.7) if IS_LAND else Cm(21)
    _ph  = Cm(21)   if IS_LAND else Cm(29.7)
    _lm  = Cm(2.0)  if IS_LAND else Cm(2.5)
    _rm  = Cm(2.0)  if IS_LAND else Cm(2.5)
    _tm  = Cm(2.0)  if IS_LAND else Cm(2.5)
    _bm  = Cm(1.5)  if IS_LAND else Cm(2.0)
    # 콘텐츠 폭 기준 차트 너비
    CW_FULL = 23 if IS_LAND else 15   # 전폭 차트 (cm)
    CW_HALF = 20 if IS_LAND else 13   # 반폭 차트 (cm)

    doc = Document()
    for sec in doc.sections:
        sec.page_width    = _pw;  sec.page_height   = _ph
        sec.left_margin   = _lm;  sec.right_margin  = _rm
        sec.top_margin    = _tm;  sec.bottom_margin = _bm
        if IS_LAND:
            # 가로형: landscape 플래그 설정
            from docx.oxml.ns import qn as _qn
            from docx.oxml import OxmlElement as _OE
            pgSz = sec._sectPr.find(_qn('w:pgSz'))
            if pgSz is None:
                pgSz = _OE('w:pgSz')
                sec._sectPr.append(pgSz)
            pgSz.set(_qn('w:orient'), 'landscape')

    # ── 제목 (FITI CI 헤더) ───────────────────────────────────────
    _logo_header(doc, "제품 불량률 분석 보고서", f"검사 기간 : {period}")

    # ── 섹션1: 전체 불량률 요약 ───────────────────────────────────
    if IS_ALL:
        _sec_title(doc, "1. 전체 불량률 요약  (1차 / 최종 / 수정합격률)")
        _hdrs1 = ["총 검사수량", "1차 불량수량", "1차 불량률", "최종 불량수량", "최종 불량률", "수정 합격률"]
        _vals1 = [f"{summary['inspec']:,} 개",
                  f"{summary['defect']:,} 개", f"{avg:.2f} %",
                  f"{summary['final_defect']:,} 개", f"{avg_final:.2f} %",
                  f"{avg_corr:.2f} %"]
        summary_tbl = doc.add_table(rows=2, cols=6)
        summary_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell, h in zip(summary_tbl.rows[0].cells, _hdrs1):
            _ct(cell, h, bold=True, sz=8.5, color=RGBColor(0xFF, 0xFF, 0xFF))
            _shd(cell, C_THEAD); _bdr(cell)
        for i, (cell, v) in enumerate(zip(summary_tbl.rows[1].cells, _vals1)):
            _ct(cell, v, bold=True, sz=12,
                color=(C_RED if i in (2, 4) else (C_GREEN if i == 5 else C_DARK)))
            _shd(cell, "F7FBFF"); _bdr(cell)
    elif IS_DUAL:
        _sec_title(doc, "1. 전체 불량률 요약  (1차 / 최종 구분)")
        _hdrs1 = ["총 검사수량", "1차 불량수량", "1차 불량률", "최종 불량수량", "최종 불량률"]
        _vals1 = [f"{summary['inspec']:,} 개",
                  f"{summary['defect']:,} 개", f"{avg:.2f} %",
                  f"{summary['final_defect']:,} 개", f"{avg_final:.2f} %"]
        summary_tbl = doc.add_table(rows=2, cols=5)
        summary_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell, h in zip(summary_tbl.rows[0].cells, _hdrs1):
            _ct(cell, h, bold=True, sz=9, color=RGBColor(0xFF, 0xFF, 0xFF))
            _shd(cell, C_THEAD); _bdr(cell)
        for i, (cell, v) in enumerate(zip(summary_tbl.rows[1].cells, _vals1)):
            _ct(cell, v, bold=True, sz=12,
                color=(C_RED if i in (2, 4) else C_DARK))
            _shd(cell, "F7FBFF"); _bdr(cell)
    else:
        _sec_title(doc, "1. 전체 불량률 요약")
        _lbl_d = "1차 불량수량" if IS_1ST else "총 불량수량"
        _lbl_r = "1차 불량률"   if IS_1ST else "전체 불량률"
        summary_tbl = doc.add_table(rows=2, cols=3)
        summary_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for cell, h in zip(summary_tbl.rows[0].cells, ["총 검사수량", _lbl_d, _lbl_r]):
            _ct(cell, h, bold=True, sz=9, color=RGBColor(0xFF, 0xFF, 0xFF))
            _shd(cell, C_THEAD); _bdr(cell)
        vals = [f"{summary['inspec']:,} 개", f"{summary['defect']:,} 개", f"{avg:.2f} %"]
        for i, (cell, v) in enumerate(zip(summary_tbl.rows[1].cells, vals)):
            _ct(cell, v, bold=True, sz=14, color=(C_RED if i == 2 else C_DARK))
            _shd(cell, "F7FBFF"); _bdr(cell)

    if MPL_OK and data["monthly"]:
        _spacer(doc)
        if IS_ALL:
            _img_para(doc, _chart_monthly_all(data["monthly"], avg, avg_final, avg_corr), CW_FULL)
        elif IS_DUAL:
            _img_para(doc, _chart_monthly_dual(data["monthly"], avg, avg_final), CW_FULL)
        else:
            _img_para(doc, _chart_monthly(data["monthly"], avg), CW_FULL)

    # ── 섹션2: 국가별 불량률 ─────────────────────────────────────
    _sec_title(doc, "2. 국가별 불량률")
    rows2c = []
    if IS_ANY_DUAL:
        for d in data["by_country"]:
            d1 = d["rate"] - avg;  d2 = d.get("final_rate", 0) - avg_final
            s1 = f"▲ +{d1:.2f}%" if d1 > 0 else f"▼ {d1:.2f}%"
            s2 = f"▲ +{d2:.2f}%" if d2 > 0 else f"▼ {d2:.2f}%"
            row_c = [
                (d["name"], False, None, WD_ALIGN_PARAGRAPH.LEFT),
                (f"{d['inspec']:,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{d['defect']:,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{d['rate']:.2f}%", False,
                 C_RED if d["rate"] > avg else None, WD_ALIGN_PARAGRAPH.CENTER),
                (s1, False, C_RED if d1 > 0 else C_GREEN, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{d.get('final_defect', 0):,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{d.get('final_rate', 0):.2f}%", False,
                 C_RED if d.get('final_rate', 0) > avg_final else None, WD_ALIGN_PARAGRAPH.CENTER),
                (s2, False, C_RED if d2 > 0 else C_GREEN, WD_ALIGN_PARAGRAPH.CENTER),
            ]
            if IS_ALL:
                row_c.append((f"{d.get('correction_rate', 0):.1f}%", False,
                    C_GREEN if d.get('correction_rate', 0) >= 80 else None, WD_ALIGN_PARAGRAPH.CENTER))
            rows2c.append(row_c)
        _hdrs2 = ["국가명","검사수량(개)","1차불량수량","1차불량률(%)","1차평균대비",
                  "최종불량수량","최종불량률(%)","최종평균대비"]
        if IS_ALL: _hdrs2.append("수정합격률(%)")
        _make_table(doc, _hdrs2, rows2c)
        if MPL_OK and data["by_country"]:
            _spacer(doc)
            if IS_ALL:
                _img_para(doc, _chart_country_all(data["by_country"], avg, avg_final, avg_corr), CW_HALF)
            else:
                _img_para(doc, _chart_country_dual(data["by_country"], avg, avg_final), CW_HALF)
    else:
        _lbl_d = "1차불량수량(개)" if IS_1ST else "불량수량(개)"
        _lbl_r = "1차불량률(%)"   if IS_1ST else "불량률(%)"
        for d in data["by_country"]:
            diff = d["rate"] - avg
            ds = f"▲ +{diff:.2f}%" if diff > 0 else f"▼ {diff:.2f}%"
            dc = C_RED if diff > 0 else C_GREEN
            rows2c.append([
                (d["name"], False, None, WD_ALIGN_PARAGRAPH.LEFT),
                (f"{d['inspec']:,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{d['defect']:,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{d['rate']:.2f}%", False,
                 C_RED if d["rate"] > avg else None, WD_ALIGN_PARAGRAPH.CENTER),
                (ds, False, dc, WD_ALIGN_PARAGRAPH.CENTER),
            ])
        _make_table(doc, ["국가명","검사수량(개)", _lbl_d, _lbl_r,"평균 대비"], rows2c)
        if MPL_OK and data["by_country"]:
            _spacer(doc)
            _img_para(doc, _chart_country(data["by_country"], avg), CW_HALF)

    # ── 섹션3: 업체별 불량률 ─────────────────────────────────────
    _sec_title(doc, "3. 업체별 불량률")
    rows3 = []
    if IS_ANY_DUAL:
        for d in data["by_client"]:
            d1 = d["rate"] - avg;  d2 = d.get("final_rate", 0) - avg_final
            s1 = f"▲ +{d1:.2f}%" if d1 > 0 else f"▼ {d1:.2f}%"
            s2 = f"▲ +{d2:.2f}%" if d2 > 0 else f"▼ {d2:.2f}%"
            row_c = [
                (d["name"], False, None, WD_ALIGN_PARAGRAPH.LEFT),
                (f"{d['inspec']:,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{d['defect']:,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{d['rate']:.2f}%", False,
                 C_RED if d["rate"] > avg else None, WD_ALIGN_PARAGRAPH.CENTER),
                (s1, False, C_RED if d1 > 0 else C_GREEN, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{d.get('final_defect', 0):,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{d.get('final_rate', 0):.2f}%", False,
                 C_RED if d.get('final_rate', 0) > avg_final else None, WD_ALIGN_PARAGRAPH.CENTER),
                (s2, False, C_RED if d2 > 0 else C_GREEN, WD_ALIGN_PARAGRAPH.CENTER),
            ]
            if IS_ALL:
                row_c.append((f"{d.get('correction_rate', 0):.1f}%", False,
                    C_GREEN if d.get('correction_rate', 0) >= 80 else None, WD_ALIGN_PARAGRAPH.CENTER))
            rows3.append(row_c)
        _hdrs3 = ["업체명","검사수량(개)","1차불량수량","1차불량률(%)","1차평균대비",
                  "최종불량수량","최종불량률(%)","최종평균대비"]
        if IS_ALL: _hdrs3.append("수정합격률(%)")
        _make_table(doc, _hdrs3, rows3)
        if MPL_OK and data["by_client"]:
            _spacer(doc)
            if IS_ALL:
                _img_para(doc, _chart_client_all(data["by_client"], avg, avg_final, avg_corr), CW_HALF)
            else:
                _img_para(doc, _chart_client_dual(data["by_client"], avg, avg_final), CW_HALF)
    else:
        _lbl_d = "1차불량수량(개)" if IS_1ST else "불량수량(개)"
        _lbl_r = "1차불량률(%)"   if IS_1ST else "불량률(%)"
        for d in data["by_client"]:
            diff = d["rate"] - avg
            ds = f"▲ +{diff:.2f}%" if diff > 0 else f"▼ {diff:.2f}%"
            dc = C_RED if diff > 0 else C_GREEN
            rows3.append([
                (d["name"], False, None, WD_ALIGN_PARAGRAPH.LEFT),
                (f"{d['inspec']:,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{d['defect']:,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{d['rate']:.2f}%", False,
                 C_RED if d["rate"] > avg else None, WD_ALIGN_PARAGRAPH.CENTER),
                (ds, False, dc, WD_ALIGN_PARAGRAPH.CENTER),
            ])
        _make_table(doc, ["업체명","검사수량(개)", _lbl_d, _lbl_r,"평균 대비"], rows3)
        if MPL_OK and data["by_client"]:
            _spacer(doc)
            _img_para(doc, _chart_client(data["by_client"], avg), CW_HALF)

    # ── 섹션4~8: 1차 불량률 기준 ─────────────────────────────────
    _note = "※ 1차 불량률 기준" if IS_ANY_DUAL else None
    _sec_title(doc, "4. 품목 유형별 현황 (의류 / 잡화 / 신발)", note=_note)
    rows4t = []
    for d in data.get("by_item_type", []):
        rows4t.append([
            (d["name"], True, None, WD_ALIGN_PARAGRAPH.LEFT),
            (f"{d['inspec']:,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
            (f"{d['defect']:,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
            (f"{d['rate']:.2f}%", False,
             C_RED if d["rate"] > avg else None, WD_ALIGN_PARAGRAPH.CENTER),
        ])
    if rows4t:
        _make_table(doc, ["품목 유형","검사수량(개)","불량수량(개)","불량률(%)"], rows4t)
    if MPL_OK and data.get("by_item_type"):
        _spacer(doc)
        img = _chart_item_type(data["by_item_type"])
        if img:
            _img_para(doc, img, CW_HALF + 1)

    # ── 섹션5: 전체 세부 불량 유형 ───────────────────────────────
    _sec_title(doc, f"5. 전체 세부 불량 유형  ※ 상위 {TOP_N}개 + 그외", note=_note)
    cum = 0
    rows3 = []
    for d in data["defect_types"]:
        cum += d["pct"]
        rows3.append([
            (d["type"], False, None, WD_ALIGN_PARAGRAPH.LEFT),
            (str(d["count"]), False, None, WD_ALIGN_PARAGRAPH.CENTER),
            (f"{d['pct']:.1f}%", False, None, WD_ALIGN_PARAGRAPH.CENTER),
            (f"{cum:.1f}%", False, None, WD_ALIGN_PARAGRAPH.CENTER),
        ])
    _make_table(doc, ["불량 유형","건수","비율(%)","누적비율(%)"], rows3)
    if MPL_OK and data["defect_types"]:
        _spacer(doc)
        _img_para(doc, _chart_defect(data["defect_types"]), CW_FULL)

    # ── 섹션6: 업체별 세부 불량 유형 ─────────────────────────────
    _sec_title(doc, "6. 업체별 세부 불량 유형", note=_note)
    type_cols = data["top_type_names"]
    rows4 = []
    for client in [d["name"] for d in data["by_client"]]:
        dmap = data["client_defect_table"].get(client, {})
        row = [(client, False, None, WD_ALIGN_PARAGRAPH.LEFT)]
        row += [(str(dmap.get(t, 0)), False, None, WD_ALIGN_PARAGRAPH.CENTER)
                for t in type_cols]
        rows4.append(row)
    _make_table(doc, ["업체명"] + type_cols, rows4)

    # ── 섹션7: 공장별 불량률 ─────────────────────────────────────
    _sec_title(doc, "7. 공장별 불량률", note=_note)
    rows5 = []
    for d in data["by_factory"]:
        rows5.append([
            (d["name"], False, None, WD_ALIGN_PARAGRAPH.LEFT),
            (f"{d['inspec']:,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
            (f"{d['defect']:,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
            (f"{d['rate']:.2f}%", False,
             C_RED if d["rate"] > avg else None, WD_ALIGN_PARAGRAPH.CENTER),
        ])
    _make_table(doc, ["공장명","검사수량(개)","불량수량(개)","불량률(%)"], rows5)
    if MPL_OK and data["by_factory"]:
        _spacer(doc)
        _img_para(doc, _chart_factory(data["by_factory"], avg), CW_FULL)

    # ── 섹션8: 공장별 세부 불량 유형 ─────────────────────────────
    _sec_title(doc, "8. 공장별 세부 불량 유형", note=_note)
    rows6 = []
    for factory in [d["name"] for d in data["by_factory"]]:
        dmap = data["factory_defect_table"].get(factory, {})
        row = [(factory, False, None, WD_ALIGN_PARAGRAPH.LEFT)]
        row += [(str(dmap.get(t, 0)), False, None, WD_ALIGN_PARAGRAPH.CENTER)
                for t in type_cols]
        rows6.append(row)
    _make_table(doc, ["공장명"] + type_cols, rows6)

    # ── 생성일 ───────────────────────────────────────────────────
    p_f = doc.add_paragraph(); p_f.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_f.paragraph_format.space_before = Pt(10)
    _run(p_f, f"보고서 생성일 : {datetime.now().strftime('%Y-%m-%d')}",
         sz=8.5, color=RGBColor(0x88, 0x88, 0x88))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ── 공장별 Word 보고서 (PDF와 동일 구조) ─────────────────────────

def _fw_chart_trend(monthly, factory, has_dual):
    """PDF와 동일한 스타일의 월별 추이 차트 반환 (BytesIO)"""
    if not MPL_OK or not monthly:
        return None
    months = [m["month"] for m in monthly if m.get("month")]
    rates1 = [m.get("rate") or 0 for m in monthly]
    if not months:
        return None

    if has_dual:
        rates2 = [m.get("final_rate") or 0 for m in monthly]
        corr   = [m.get("correction_rate") or 0 for m in monthly]
        has_final = any(r > 0 for r in rates2)

        if has_final:
            fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(9, 5.5), sharex=True,
                                            gridspec_kw={'height_ratios': [3, 2]})
        else:
            fig, ax1 = plt.subplots(figsize=(9, 3.5))
            ax2 = None

        ax1.plot(months, rates1, marker='o', lw=2.5, color='#2B5BA8',
                 label='1차 불량률', ms=6, markerfacecolor='white', markeredgewidth=2)
        if has_final:
            ax1.plot(months, rates2, marker='s', lw=2.5, color='#C0392B',
                     label='최종 불량률', ms=6, markerfacecolor='white', markeredgewidth=2)
        ax1.set_title(f"{factory} — 월별 불량률 추이 (1차 vs 최종)", fontsize=12, pad=10)
        ax1.set_ylabel("불량률 (%)"); ax1.set_ylim(bottom=0)
        ax1.legend(fontsize=9, loc='upper right')
        ax1.grid(axis='y', linestyle='--', alpha=0.4)

        if ax2 is not None and any(c > 0 for c in corr):
            _colors = ['#27AE60' if c >= 80 else ('#E67E22' if c >= 60 else '#C0392B') for c in corr]
            ax2.bar(months, corr, color=_colors, width=0.5, alpha=0.85)
            for i, (m, c) in enumerate(zip(months, corr)):
                if c > 0:
                    ax2.text(i, c + 1, f'{c:.0f}%', ha='center', va='bottom', fontsize=8)
            ax2.set_ylim(0, 110); ax2.set_ylabel("수정합격률 (%)")
            ax2.set_title("수정 합격률 추이", fontsize=10, pad=6)
            ax2.grid(axis='y', linestyle='--', alpha=0.3)
            ax2.axhline(80, color='#27AE60', lw=1, linestyle=':', alpha=0.7)
    else:
        fig, ax1 = plt.subplots(figsize=(9, 3.5))
        if len(months) == 1:
            ax1.bar(months, rates1, color='#2B5BA8', width=0.4)
        else:
            ax1.plot(months, rates1, marker='o', lw=2.5, color='#2B5BA8',
                     ms=6, markerfacecolor='white', markeredgewidth=2)
            ax1.fill_between(months, rates1, alpha=0.08, color='#2B5BA8')
        ax1.set_title(f"{factory} — 월별 1차 불량률 추이", fontsize=13, pad=12)
        ax1.set_ylabel("불량률 (%)"); ax1.set_ylim(bottom=0)
        ax1.grid(axis='y', linestyle='--', alpha=0.4)

    plt.xticks(rotation=30, ha='right', fontsize=9)
    fig.tight_layout()
    return _fig_bytes(fig)


def _fw_chart_defect(top5, factory):
    """PDF와 동일한 스타일의 불량 유형 바차트"""
    if not MPL_OK or not top5:
        return None
    names = [d["name"] for d in reversed(top5)]
    qtys  = [d["qty"]  for d in reversed(top5)]
    bar_colors = ['#2B5BA8','#4472C4','#5B9BD5','#70AD47','#ED7D31',
                  '#8E44AD','#16A085'][::-1]
    fig, ax = plt.subplots(figsize=(9, 3))
    bars = ax.barh(names, qtys, color=bar_colors[:len(names)], edgecolor='white')
    for bar, qty in zip(bars, qtys):
        ax.text(bar.get_width() + max(qtys)*0.01, bar.get_y() + bar.get_height()/2,
                f'{qty:,}개', va='center', fontsize=9)
    ax.set_title(f"주요 불량 유형 TOP {len(top5)}", fontsize=13, pad=12)
    ax.set_xlabel("불량 수량 (개)")
    ax.grid(axis='x', linestyle='--', alpha=0.3)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    fig.tight_layout()
    return _fig_bytes(fig)


def generate_factory_word(detail: dict) -> bytes:
    """
    PDF와 동일한 구조의 공장별 .docx 반환
    기본정보 → KPI → 월별추이 → TOP5불량유형 → 월별검사실적
    """
    if not DOCX_OK:
        raise RuntimeError("python-docx 설치 필요")

    factory      = detail.get('factory', '공장')
    region1      = detail.get('region1', '')
    region2      = detail.get('region2', '')
    buyers       = ", ".join(detail.get('buyers', [])) or "—"
    record_count = detail.get('record_count', 0)
    total_inspec = detail.get('total_inspec', 0)
    total_defect = detail.get('total_defect', 0)
    total_final  = detail.get('total_final_defect', 0)
    total_second = detail.get('total_second_inspec', 0)
    avg_rate     = detail.get('avg_rate') or 0.0
    final_rate   = detail.get('final_rate') or 0.0
    corr_rate    = detail.get('correction_rate') or 0.0
    has_dual     = (total_final > 0 or total_second > 0)
    monthly      = detail.get('monthly', [])
    top5         = detail.get('top7_defects', [])
    today        = datetime.now().strftime('%Y년 %m월 %d일')

    doc = Document()
    for sec in doc.sections:
        sec.page_width    = Cm(21);   sec.page_height   = Cm(29.7)
        sec.left_margin   = Cm(2.0);  sec.right_margin  = Cm(2.0)
        sec.top_margin    = Cm(2.0);  sec.bottom_margin = Cm(2.0)

    # ── 헤더 ─────────────────────────────────────────────────────
    _hr(doc)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    _run(p, "공장 불량률 분석 보고서", sz=22, bold=True, color=C_HEADER)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(2)
    _run(p2, factory, sz=16, bold=True, color=C_SUB)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0); p3.paragraph_format.space_after = Pt(6)
    _run(p3, f"보고일: {today}", sz=10, color=RGBColor(0x6C,0x75,0x7D))
    _hr(doc)

    # ── 기본 정보 ─────────────────────────────────────────────────
    _sec_title(doc, "기본 정보")
    info_tbl = doc.add_table(rows=2, cols=4)
    info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ["지역", f"{region1} {region2}".strip(), "바이어", buyers],
        ["검사 건수", f"{record_count:,}건", "총 검사수량", f"{total_inspec:,}개"],
    ]
    for ri, row_data in enumerate(info_data):
        for ci, val in enumerate(row_data):
            cell = info_tbl.rows[ri].cells[ci]
            is_label = (ci % 2 == 0)
            _ct(cell, val, bold=is_label, sz=9.5,
                color=C_SUB if is_label else C_DARK,
                align=WD_ALIGN_PARAGRAPH.LEFT)
            _shd(cell, "E8F0FE" if is_label else "FFFFFF")
            _bdr(cell)

    _spacer(doc)

    # ── KPI 카드 ─────────────────────────────────────────────────
    def _rate_color(r, is_corr=False):
        if is_corr:
            return C_GREEN if r >= 80 else (RGBColor(0xE6,0x7E,0x22) if r >= 60 else C_RED)
        return C_GREEN if r < 5 else (RGBColor(0xE6,0x7E,0x22) if r < 10 else C_RED)

    if has_dual:
        kpi_hdrs = ["1차 불량률", "최종 불량률", "수정 합격률", "1차 불량수량", "분석 기간"]
        kpi_vals = [f"{avg_rate:.2f}%", f"{final_rate:.2f}%", f"{corr_rate:.1f}%",
                    f"{total_defect:,}개", f"{len(monthly)}개월"]
        kpi_colors = [_rate_color(avg_rate), _rate_color(final_rate),
                      _rate_color(corr_rate, True), C_HEADER, RGBColor(0x6C,0x75,0x7D)]
        kpi_tbl = doc.add_table(rows=2, cols=5)
    else:
        kpi_hdrs = ["평균 불량률", "총 불량수량", "분석 기간"]
        kpi_vals = [f"{avg_rate:.2f}%", f"{total_defect:,}개", f"{len(monthly)}개월"]
        kpi_colors = [_rate_color(avg_rate), C_HEADER, RGBColor(0x6C,0x75,0x7D)]
        kpi_tbl = doc.add_table(rows=2, cols=3)

    kpi_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, h in zip(kpi_tbl.rows[0].cells, kpi_hdrs):
        _ct(cell, h, bold=True, sz=8.5, color=RGBColor(0xFF,0xFF,0xFF))
        _shd(cell, C_THEAD); _bdr(cell)
    for cell, v, col in zip(kpi_tbl.rows[1].cells, kpi_vals, kpi_colors):
        _ct(cell, v, bold=True, sz=14, color=col)
        _shd(cell, "E8F0FE"); _bdr(cell)

    _spacer(doc)

    # ── 월별 불량률 추이 ──────────────────────────────────────────
    _sec_title(doc, "월별 불량률 추이")
    if MPL_OK and monthly:
        trend_img = _fw_chart_trend(monthly, factory, has_dual)
        if trend_img:
            _img_para(doc, trend_img, 15)
    _spacer(doc)

    # ── 주요 불량 유형 ────────────────────────────────────────────
    if top5:
        _sec_title(doc, f"주요 불량 유형 TOP {len(top5)}")
        defect_img = _fw_chart_defect(top5, factory)
        if defect_img:
            _img_para(doc, defect_img, 15)
        _spacer(doc)
        d_rows = []
        for i, d in enumerate(top5):
            d_rows.append([
                (f"{i+1}위", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (d["name"], False, None, WD_ALIGN_PARAGRAPH.LEFT),
                (f"{d['qty']:,}개", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{d['pct']}%", False, None, WD_ALIGN_PARAGRAPH.CENTER),
            ])
        _make_table(doc, ["순위", "표준 불량명", "불량 수량", "비율"], d_rows)
        _spacer(doc)

    # ── 월별 검사 실적 ────────────────────────────────────────────
    _sec_title(doc, "월별 검사 실적")
    if has_dual:
        m_rows = []
        for m in monthly:
            r1 = m.get("rate"); r2 = m.get("final_rate"); rc = m.get("correction_rate")
            m_rows.append([
                (m["month"], False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{m.get('inspec',0):,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{m.get('defect',0):,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{r1:.2f}%" if r1 is not None else "—", False,
                 _rate_color(r1) if r1 is not None else None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{m.get('final_defect',0):,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{r2:.2f}%" if r2 is not None else "—", False,
                 _rate_color(r2) if r2 is not None else None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{rc:.1f}%" if rc is not None else "—", False,
                 _rate_color(rc, True) if rc is not None else None, WD_ALIGN_PARAGRAPH.CENTER),
            ])
        _make_table(doc,
            ["연월","검사수량","1차불량수량","1차불량률","최종불량수량","최종불량률","수정합격률"],
            m_rows)
    else:
        m_rows = []
        for m in monthly:
            rate = m.get("rate")
            if rate is None:
                rate_str, eval_str = "—", "—"
            else:
                rate_str = f"{rate:.2f}%"
                eval_str = "우수" if rate < 5.0 else ("양호" if rate < 7.0 else ("주의" if rate < 10.0 else "불량"))
            m_rows.append([
                (m["month"], False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{m.get('inspec',0):,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (f"{m.get('defect',0):,}", False, None, WD_ALIGN_PARAGRAPH.CENTER),
                (rate_str, False, _rate_color(rate) if rate is not None else None, WD_ALIGN_PARAGRAPH.CENTER),
                (eval_str, False, None, WD_ALIGN_PARAGRAPH.CENTER),
            ])
        _make_table(doc, ["연월","검사수량","불량수량","불량률","평가"], m_rows)

    # ── 푸터 ─────────────────────────────────────────────────────
    _hr(doc)
    p_f = doc.add_paragraph(); p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_f.paragraph_format.space_before = Pt(4)
    _run(p_f,
         f"본 보고서는 FITI 제품평가팀 검사 데이터 기반으로 자동 생성되었습니다.  |  {today}",
         sz=8, color=RGBColor(0x6C,0x75,0x7D))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
